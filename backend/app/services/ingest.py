import os
import mimetypes
import logging
from pathlib import Path
from fastapi import UploadFile
from sqlalchemy.orm import Session
from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from app import db

from ..settings import settings
from .. import models
from .chunking import chunk_text

logger = logging.getLogger(__name__)
# Try to use improved chunking if available
try:
    from .chunking_improved import chunk_text_improved
    USE_IMPROVED_CHUNKING = True
except ImportError:
    USE_IMPROVED_CHUNKING = False

from .embeddings import embed_text
from .pinecone_client import index as pinecone_index
from .r2_storage import get_r2_storage
from .storage_utils import generate_storage_key, get_scope_from_category, sanitize_filename
from .meeting_summary import generate_meeting_summary

def _ensure_dir(path: str) -> None:
    Path(path).mkdir(parents=True, exist_ok=True)


def _save_upload(customer_id: str, upload_file: UploadFile) -> str:
    base_dir = settings.UPLOAD_DIR
    customer_dir = os.path.join(base_dir, customer_id)
    _ensure_dir(customer_dir)

    file_path = os.path.join(customer_dir, upload_file.filename)
    with open(file_path, "wb") as f:
        f.write(upload_file.file.read())
    return file_path


def _extract_text(file_path: str) -> tuple[str, int | None]:
    """
    Extract text from file and return (text, page_count).
    For PDFs, page_count is the number of pages; for others, None.
    """
    ext = Path(file_path).suffix.lower()
    page_count = None

    if ext == ".docx":
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return text.strip(), page_count

    if ext == ".pdf":
        reader = PdfReader(file_path)
        pages = []
        page_count = len(reader.pages)
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip(), page_count

    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()
        return text, page_count
    except Exception:
        return "", page_count


def ingest_document(db: Session, customer_id: str, doc_type: str, upload_file: UploadFile, project_id: str | None = None, document_category: str = "project") -> models.Document:
    
    # Read file content
    file_content = upload_file.file.read()
    upload_file.file.seek(0)  # Reset for potential reuse
    
    # Get file metadata
    file_size = len(file_content)
    mime_type, _ = mimetypes.guess_type(upload_file.filename)
    if not mime_type:
        # Fallback to common types
        ext = Path(upload_file.filename).suffix.lower()
        mime_type_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain",
        }
        mime_type = mime_type_map.get(ext, "application/octet-stream")

    # 2) Create document row first (to get document_id for storage key)
    doc = models.Document(
        customer_id=customer_id,
        project_id=project_id,
        document_category=document_category,
        doc_type=doc_type,
        filename=upload_file.filename,
        storage_provider=settings.STORAGE_PROVIDER or "r2",
        file_size=file_size,
        mime_type=mime_type,
        processing_status="processing",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # 3) Upload to R2 (if configured) or save locally
    storage_key = None
    file_path = None
    
    if settings.STORAGE_PROVIDER == "r2":
        try:
            r2 = get_r2_storage()
            if r2._is_configured():
                # Generate storage key
                scope = get_scope_from_category(document_category)
                storage_key = generate_storage_key(
                    customer_id=customer_id,
                    project_id=project_id,
                    scope=scope,
                    doc_type=doc_type,
                    document_id=doc.id,
                    filename=upload_file.filename
                )
                
                # Upload to R2
                r2.upload_file(
                    file_bytes=file_content,
                    key=storage_key,
                    content_type=mime_type
                )
                
                # Update document with storage_key
                doc.storage_key = storage_key
                db.commit()
            else:
                # Fallback to local storage if R2 not configured
                file_path = _save_upload(customer_id, upload_file)
                doc.storage_path = file_path
                doc.storage_provider = "local"
                db.commit()
        except Exception as e:
            # Fallback to local storage on error
            logger.warning(f"R2 upload failed, falling back to local storage: {e}")
            file_path = _save_upload(customer_id, upload_file)
            doc.storage_path = file_path
            doc.storage_provider = "local"
            db.commit()
    else:
        # Local storage
        file_path = _save_upload(customer_id, upload_file)
        doc.storage_path = file_path
        doc.storage_provider = "local"
        db.commit()

    # 4) Extract text (from file_path if local, or from bytes if R2)
    if file_path and os.path.exists(file_path):
        extracted, page_count = _extract_text(file_path)
    else:
        # Extract from bytes for R2
        # Save temporarily to extract text
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(upload_file.filename).suffix) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            extracted, page_count = _extract_text(tmp_path)
        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except:
                pass
    if not extracted:
        extracted = "(No text extracted from this file. Try a .txt/.docx/.pdf with selectable text.)"

    # 4) Generate summary for meeting minutes (for RAG efficiency)
    summary = None
    if doc_type == "meeting_minutes" and len(extracted.strip()) > 200:
        try:
            summary = generate_meeting_summary(extracted)
            logger.info(f"Generated summary for meeting minutes document {doc.id}")
        except Exception as e:
            logger.warning(f"Failed to generate meeting summary: {e}, using full text")
            summary = None

    # Update document with page_count and set status to "completed"
    doc.page_count = page_count
    doc.processing_status = "completed"
    db.commit()

    # 5) Store extracted text and summary
    doc_text = models.DocumentText(
        document_id=doc.id,
        extracted_text=extracted,
        summary=summary
    )
    db.add(doc_text)
    db.commit()

    # 5) Get customer name for metadata enrichment
    customer = db.query(models.Customer).filter(models.Customer.id == customer_id).first()
    customer_name = customer.name if customer else None
    
    # Only index project documents in Pinecone (skip customer documents)
    if document_category == "project":
        # 6) For meeting minutes, use summary instead of full text for RAG indexing
        text_to_index = summary if (doc_type == "meeting_minutes" and summary) else extracted
        
        # Chunk and store chunks (use improved chunking if available)
        if USE_IMPROVED_CHUNKING:
            chunks = chunk_text_improved(text_to_index, chunk_size=900, overlap=150)
        else:
            chunks = chunk_text(text_to_index, chunk_size=900, overlap=150)
        
        for i, ch in enumerate(chunks):
            vector = embed_text(ch)

            vector_id = f"{doc.id}_{i}"
            namespace = (settings.PINECONE_NAMESPACE or "").strip()

            # Build metadata with enrichment (customer_name, document_filename)
            metadata = {
                "customer_id": customer_id,
                "document_id": doc.id,
                "chunk_index": i,
                "doc_type": doc_type,
                "document_category": "project",  # Always "project" for indexed documents
                "uploaded_at": doc.uploaded_at.isoformat(),
                "text": ch,
            }
            
            # Add enriched metadata
            if customer_name:
                metadata["customer_name"] = customer_name
            if doc.filename:
                metadata["document_filename"] = doc.filename
            
            # Only include project_id if it's not None
            if project_id is not None:
                metadata["project_id"] = project_id

            pinecone_index.upsert(
                vectors=[
                    {
                        "id": vector_id,
                        "values": vector,
                        "metadata": metadata,
                    }
                ],
                # ✅ This is what creates "context" namespace
                namespace=namespace if namespace else None,
            )

            row = models.Chunk(
                document_id=doc.id,
                chunk_index=i,
                chunk_text=ch,
                pinecone_vector_id=vector_id,
            )
            db.add(row)

        db.commit()
    else:
        # Customer documents are not indexed in Pinecone
        # Still store the document in the database, but no chunks/embeddings
        pass
    

    return doc
